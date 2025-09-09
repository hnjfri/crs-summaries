#!/usr/bin/env python3
"""
CRS Summary Extractor

This script fetches Congressional Research Service (CRS) documents from the Congress.gov API
and creates a Word document with formatted reports including titles, authors, dates, and summaries.

Architecture Decision: Single-Class Design with Comprehensive Error Handling

Decision: Use a single CRSSummaryExtractor class that encapsulates all functionality
with comprehensive error handling, input validation, and structured logging.

Context: This is a focused script for a specific task (CRS data extraction).
While we could split into multiple classes, the operations are tightly coupled
and the script's scope is well-defined.

Trade-offs:
- Pro: Simple, focused, easy to understand and maintain
- Pro: All related functionality in one place
- Con: Single class handles multiple responsibilities
- Mitigation: Clear method separation and comprehensive documentation

Requirements:
- Only includes documents with "status": "Active"
- For duplicate id+publishDate combinations, keeps only the highest version number
- Outputs Word document formatted for professional presentation
- Comprehensive error handling and input validation
- Structured logging with security-conscious practices
"""

import hashlib
import json
import logging
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

# Configure logging first - centralized logging setup
def configure_logging(verbose: bool = False, json_logs: bool = False) -> logging.Logger:
    """Centralized logger configuration following enterprise standards.
    
    Args:
        verbose: Enable debug-level logging
        json_logs: Output structured JSON logs instead of human-readable format
        
    Returns:
        Configured logger instance for the CRS extractor
    """
    logger = logging.getLogger('crs_extractor')
    logger.setLevel(logging.DEBUG if verbose else logging.INFO)
    logger.handlers.clear()
    logger.propagate = False

    console = logging.StreamHandler(sys.stdout)
    console.setLevel(logging.WARNING)  # Only warnings and errors to console
    
    if json_logs:
        class JsonFormatter(logging.Formatter):
            def format(self, record: logging.LogRecord) -> str:
                payload = {
                    'timestamp': self.formatTime(record, '%Y-%m-%d %H:%M:%S'),
                    'logger': record.name,
                    'level': record.levelname,
                    'message': record.getMessage(),
                    'module': record.module,
                    'function': record.funcName,
                    'line': record.lineno,
                }
                if hasattr(record, 'operation'):
                    payload['operation'] = record.operation
                if hasattr(record, 'context'):
                    payload['context'] = record.context
                return json.dumps(payload)
        console.setFormatter(JsonFormatter())
    else:
        console.setFormatter(logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            '%Y-%m-%d %H:%M:%S'
        ))
    
    logger.addHandler(console)
    return logger


# Custom exception classes for better error handling
class CRSExtractorError(Exception):
    """Base exception for CRS Summary Extractor."""
    pass


class ConfigurationError(CRSExtractorError):
    """Raised when configuration is invalid or missing."""
    pass


class ValidationError(CRSExtractorError):
    """Raised when data validation fails."""
    def __init__(self, message: str, field_name: Optional[str] = None, invalid_value: Any = None):
        super().__init__(message)
        self.field_name = field_name
        self.invalid_value = invalid_value


class APIError(CRSExtractorError):
    """Raised when API operations fail."""
    def __init__(self, message: str, status_code: Optional[int] = None, response_data: Optional[Dict[str, Any]] = None):
        super().__init__(message)
        self.status_code = status_code
        self.response_data = response_data


class DataProcessingError(CRSExtractorError):
    """Raised when data processing operations fail."""
    pass


# Load environment variables from the correct location
load_dotenv('.env')

class InputValidator:
    """Validates all external inputs to the system following security standards."""
    
    @staticmethod
    def validate_api_key(api_key: Optional[str]) -> str:
        """Validate API key format and presence.
        
        Args:
            api_key: The API key to validate
            
        Returns:
            The validated API key
            
        Raises:
            ConfigurationError: If API key is missing or invalid format
        """
        if not api_key or not isinstance(api_key, str):
            raise ConfigurationError(
                "CONGRESSGOV_API_KEY environment variable required. "
                "Please check your .env file."
            )
        
        api_key = api_key.strip()
        if len(api_key) < 10:  # Basic length check
            raise ConfigurationError("API key appears to be too short")
        
        return api_key
    
    @staticmethod
    def validate_limit_offset(limit: int, offset: int) -> Tuple[int, int]:
        """Validate API pagination parameters.
        
        Args:
            limit: Maximum number of records to fetch
            offset: Number of records to skip
            
        Returns:
            Validated (limit, offset) tuple
            
        Raises:
            ValidationError: If parameters are out of valid range
        """
        if not isinstance(limit, int) or limit < 1 or limit > 1000:
            raise ValidationError("Limit must be an integer between 1 and 1000", "limit", limit)
        
        if not isinstance(offset, int) or offset < 0:
            raise ValidationError("Offset must be a non-negative integer", "offset", offset)
        
        return limit, offset
    
    @staticmethod
    def validate_report_id(report_id: Optional[str]) -> str:
        """Validate CRS report ID format.
        
        Args:
            report_id: The report ID to validate
            
        Returns:
            The validated report ID
            
        Raises:
            ValidationError: If report ID is invalid
        """
        if not report_id or not isinstance(report_id, str):
            raise ValidationError("Report ID must be a non-empty string", "report_id", report_id)
        
        report_id = report_id.strip()
        
        # Basic format validation - CRS report IDs are typically alphanumeric with hyphens
        if not re.match(r'^[a-zA-Z0-9-]+$', report_id):
            raise ValidationError(
                "Report ID contains invalid characters. Only letters, numbers, and hyphens allowed.",
                "report_id", report_id
            )
        
        if len(report_id) > 50:  # Reasonable upper bound
            raise ValidationError("Report ID too long", "report_id", report_id)
        
        return report_id
    
    @staticmethod
    def validate_filename(filename: str) -> Path:
        """Validate output filename for security.
        
        Args:
            filename: The filename to validate
            
        Returns:
            Validated Path object
            
        Raises:
            ValidationError: If filename is unsafe
        """
        if not filename or not isinstance(filename, str):
            raise ValidationError("Filename must be a non-empty string", "filename", filename)
        
        # Convert to Path and resolve to prevent directory traversal
        try:
            path = Path(filename).resolve()
        except (OSError, ValueError) as e:
            raise ValidationError(f"Invalid filename format: {filename}") from e
        
        # Ensure it's in current directory (security measure)
        current_dir = Path.cwd().resolve()
        if not str(path).startswith(str(current_dir)):
            raise ValidationError("Filename must be in current directory", "filename", filename)
        
        # Check extension
        if path.suffix.lower() != '.docx':
            raise ValidationError("Filename must have .docx extension", "filename", filename)
        
        return path


class CRSSummaryExtractor:
    """Extracts and processes CRS summaries from Congress.gov API.
    
    This class handles the complete workflow of:
    1. Fetching CRS document lists from the API
    2. Filtering for active documents only
    3. Removing duplicates based on ID and publish date
    4. Fetching detailed information including summaries
    5. Creating formatted Word document output
    
    All operations include comprehensive error handling, input validation,
    and structured logging following enterprise security standards.
    
    Example:
        >>> extractor = CRSSummaryExtractor()
        >>> extractor.run()
        # Creates crs_summaries.docx with processed data
    """
    
    def __init__(self, verbose: bool = False, json_logs: bool = False):
        """Initialize the CRS Summary Extractor.
        
        Args:
            verbose: Enable debug-level logging
            json_logs: Output structured JSON logs
            
        Raises:
            ConfigurationError: If API key is missing or invalid
        """
        # Set up logging first
        self.logger = configure_logging(verbose, json_logs)
        
        # Generate correlation ID for this session
        self.correlation_id = hashlib.sha256(
            f"{datetime.now().isoformat()}-{os.getpid()}".encode()
        ).hexdigest()[:8]
        
        self.logger.info("Initializing CRS Summary Extractor", extra={
            'operation': 'init',
            'correlation_id': self.correlation_id
        })
        
        # Validate and set up API configuration
        raw_api_key = os.getenv('CONGRESSGOV_API_KEY')
        self.api_key = InputValidator.validate_api_key(raw_api_key)
        
        # Log successful initialization (without exposing the key)
        api_key_hash = hashlib.sha256(self.api_key.encode()).hexdigest()[:8]
        self.logger.info("API key loaded successfully", extra={
            'operation': 'init',
            'api_key_hash': api_key_hash,
            'correlation_id': self.correlation_id
        })
        
        self.base_url = "https://api.congress.gov/v3/crsreport"
        self.session = requests.Session()
        # Note: timeout is set per request, not on session
        
        # Set up rate limiting
        self.last_request_time = 0.0
        self.min_request_interval = 0.1  # 100ms between requests
    
    def _rate_limit(self) -> None:
        """Implement rate limiting to be respectful to the API."""
        current_time = time.time()
        time_since_last = current_time - self.last_request_time
        
        if time_since_last < self.min_request_interval:
            sleep_time = self.min_request_interval - time_since_last
            time.sleep(sleep_time)
        
        self.last_request_time = time.time()
    
    def _sanitize_for_logging(self, data: Any, max_length: int = 100) -> str:
        """Sanitize data for safe logging without exposing sensitive content.
        
        Args:
            data: Data to sanitize
            max_length: Maximum length of sanitized string
            
        Returns:
            Sanitized string safe for logging
        """
        if isinstance(data, dict):
            # For dictionaries, log structure but not content
            keys = list(data.keys())[:5]  # First 5 keys only
            return f"dict with keys: {keys}{'...' if len(data) > 5 else ''}"
        elif isinstance(data, list):
            return f"list with {len(data)} items"
        elif isinstance(data, str):
            if len(data) > max_length:
                # For strings, show hash instead of content
                content_hash = hashlib.sha256(data.encode()).hexdigest()[:8]
                return f"string({len(data)} chars, hash: {content_hash})"
            else:
                return f"string({len(data)} chars)"
        else:
            return f"{type(data).__name__}: {str(data)[:max_length]}"
    
    def fetch_crs_list(self, limit: int = 250, offset: int = 0) -> List[Dict[str, Any]]:
        """Fetch list of CRS documents from the API with comprehensive error handling.
        
        Args:
            limit: Maximum number of documents to fetch (1-1000)
            offset: Number of documents to skip (for pagination)
            
        Returns:
            List of CRS document dictionaries
            
        Raises:
            ValidationError: If parameters are invalid
            APIError: If API request fails
        """
        # Input validation
        limit, offset = InputValidator.validate_limit_offset(limit, offset)
        
        url = f"{self.base_url}?api_key={self.api_key}&format=json&limit={limit}&offset={offset}"
        
        # Rate limiting
        self._rate_limit()
        
        start_time = time.time()
        
        self.logger.info("Fetching CRS document list", extra={
            'operation': 'fetch_list',
            'limit': limit,
            'offset': offset,
            'correlation_id': self.correlation_id
        })
        
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            # Validate response is JSON
            try:
                data = response.json()
            except json.JSONDecodeError as e:
                raise APIError(f"Invalid JSON response from API: {e}") from e
            
            # Validate response structure
            if not isinstance(data, dict):
                raise APIError(f"Expected dict response, got {type(data)}")
            
            # Extract the reports array from the response
            reports = []
            if 'CRSReports' in data:
                reports = data['CRSReports']
            elif 'reports' in data:
                reports = data['reports']
            else:
                self.logger.warning("Unexpected response structure", extra={
                    'operation': 'fetch_list',
                    'available_keys': list(data.keys()),
                    'correlation_id': self.correlation_id
                })
                return []
            
            # Validate reports is a list
            if not isinstance(reports, list):
                raise APIError(f"Expected list of reports, got {type(reports)}")
            
            duration = time.time() - start_time
            self.logger.info("Successfully fetched CRS document list", extra={
                'operation': 'fetch_list',
                'count': len(reports),
                'duration': duration,
                'status': 'success',
                'correlation_id': self.correlation_id
            })
            
            return reports
                
        except requests.Timeout as e:
            duration = time.time() - start_time
            error_msg = f"API request timed out after {duration:.1f}s"
            self.logger.error(error_msg, extra={
                'operation': 'fetch_list',
                'error_type': 'timeout',
                'duration': duration,
                'correlation_id': self.correlation_id
            })
            raise APIError(error_msg) from e
            
        except requests.HTTPError as e:
            duration = time.time() - start_time
            status_code = e.response.status_code if e.response else None
            error_msg = f"HTTP error {status_code}: {e}"
            
            self.logger.error("API HTTP error", extra={
                'operation': 'fetch_list',
                'error_type': 'http_error',
                'status_code': status_code,
                'duration': duration,
                'correlation_id': self.correlation_id
            })
            
            # Handle specific HTTP errors
            if status_code == 401:
                raise APIError("API authentication failed. Check your API key.", status_code) from e
            elif status_code == 429:
                raise APIError("API rate limit exceeded. Please try again later.", status_code) from e
            elif status_code is not None and status_code >= 500:
                raise APIError(f"API server error: {status_code}", status_code) from e
            else:
                raise APIError(error_msg, status_code) from e
                
        except requests.RequestException as e:
            duration = time.time() - start_time
            error_msg = f"Network error fetching CRS list: {e}"
            self.logger.error("Network error", extra={
                'operation': 'fetch_list',
                'error_type': 'network_error',
                'duration': duration,
                'correlation_id': self.correlation_id
            })
            raise APIError(error_msg) from e
    
    def fetch_crs_detail(self, report_id: str) -> Optional[Dict[str, Any]]:
        """Fetch detailed information for a specific CRS document.
        
        Args:
            report_id: The ID of the CRS report to fetch
            
        Returns:
            Dictionary containing detailed report information, or None if failed
            
        Raises:
            ValidationError: If report_id is invalid
            APIError: If API request fails critically
        """
        # Input validation
        validated_id = InputValidator.validate_report_id(report_id)
        
        url = f"{self.base_url}/{validated_id}?api_key={self.api_key}&format=json"
        
        # Rate limiting
        self._rate_limit()
        
        start_time = time.time()
        
        # Hash the report ID for safe logging
        id_hash = hashlib.sha256(validated_id.encode()).hexdigest()[:8]
        
        self.logger.debug("Fetching CRS report details", extra={
            'operation': 'fetch_detail',
            'report_id_hash': id_hash,
            'correlation_id': self.correlation_id
        })
        
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            # Validate response is JSON
            try:
                data = response.json()
            except json.JSONDecodeError as e:
                self.logger.warning("Invalid JSON in detail response", extra={
                    'operation': 'fetch_detail',
                    'report_id_hash': id_hash,
                    'error': str(e),
                    'correlation_id': self.correlation_id
                })
                return None
            
            # Validate response structure
            if not isinstance(data, dict):
                self.logger.warning("Unexpected detail response type", extra={
                    'operation': 'fetch_detail',
                    'report_id_hash': id_hash,
                    'response_type': type(data).__name__,
                    'correlation_id': self.correlation_id
                })
                return None
            
            # Extract the report details
            if 'CRSReport' in data:
                detail = data['CRSReport']
                
                duration = time.time() - start_time
                self.logger.debug("Successfully fetched report details", extra={
                    'operation': 'fetch_detail',
                    'report_id_hash': id_hash,
                    'duration': duration,
                    'has_summary': 'summary' in detail,
                    'has_authors': 'authors' in detail,
                    'correlation_id': self.correlation_id
                })
                
                return detail  # type: ignore[no-any-return]
            else:
                self.logger.warning("Missing CRSReport in response", extra={
                    'operation': 'fetch_detail',
                    'report_id_hash': id_hash,
                    'available_keys': list(data.keys()),
                    'correlation_id': self.correlation_id
                })
                return None
                
        except requests.HTTPError as e:
            duration = time.time() - start_time
            status_code = e.response.status_code if e.response else None
            
            if status_code == 404:
                # 404 is expected for some report IDs, log as debug
                self.logger.debug("Report not found", extra={
                    'operation': 'fetch_detail',
                    'report_id_hash': id_hash,
                    'status_code': status_code,
                    'duration': duration,
                    'correlation_id': self.correlation_id
                })
            else:
                # Other HTTP errors are more serious
                self.logger.warning("HTTP error fetching report details", extra={
                    'operation': 'fetch_detail',
                    'report_id_hash': id_hash,
                    'status_code': status_code,
                    'duration': duration,
                    'correlation_id': self.correlation_id
                })
            
            return None
            
        except requests.RequestException as e:
            duration = time.time() - start_time
            self.logger.warning("Network error fetching report details", extra={
                'operation': 'fetch_detail',
                'report_id_hash': id_hash,
                'error_type': type(e).__name__,
                'duration': duration,
                'correlation_id': self.correlation_id
            })
            return None
    
    def filter_active_reports(self, reports: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Filter reports to only include those with status 'Active'.
        
        Args:
            reports: List of report dictionaries to filter
            
        Returns:
            List of reports with 'Active' status only
            
        Raises:
            ValidationError: If reports is not a valid list
        """
        if not isinstance(reports, list):
            raise ValidationError("Reports must be a list", "reports", type(reports))
        
        start_time = time.time()
        
        active_reports = []
        inactive_count = 0
        
        for report in reports:
            if not isinstance(report, dict):
                self.logger.warning("Skipping non-dict report", extra={
                    'operation': 'filter_active',
                    'report_type': type(report).__name__,
                    'correlation_id': self.correlation_id
                })
                continue
            
            status = report.get('status')
            if status == 'Active':
                active_reports.append(report)
            else:
                inactive_count += 1
        
        duration = time.time() - start_time
        self.logger.info("Filtered reports by status", extra={
            'operation': 'filter_active',
            'total_input': len(reports),
            'active_count': len(active_reports),
            'inactive_count': inactive_count,
            'duration': duration,
            'correlation_id': self.correlation_id
        })
        
        return active_reports
    
    def deduplicate_reports(self, reports: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Remove duplicates by keeping only the highest version number.
        
        For reports with the same ID and publish date, only the report with
        the highest version number is retained. This ensures we have the most
        current version of each document.
        
        Args:
            reports: List of report dictionaries to deduplicate
            
        Returns:
            List of deduplicated reports with highest versions only
            
        Raises:
            ValidationError: If reports is not a valid list
            DataProcessingError: If version numbers are invalid
        """
        if not isinstance(reports, list):
            raise ValidationError("Reports must be a list", "reports", type(reports))
        
        start_time = time.time()
        
        # Group by id and publishDate
        groups: Dict[Tuple[Optional[str], Optional[str]], List[Dict[str, Any]]] = {}
        invalid_reports = 0
        
        for report in reports:
            if not isinstance(report, dict):
                invalid_reports += 1
                continue
            
            report_id = report.get('id')
            publish_date = report.get('publishDate')
            
            # Skip reports without essential identifiers
            if not report_id:
                invalid_reports += 1
                self.logger.debug("Skipping report without ID", extra={
                    'operation': 'deduplicate',
                    'correlation_id': self.correlation_id
                })
                continue
            
            key = (report_id, publish_date)
            if key not in groups:
                groups[key] = []
            groups[key].append(report)
        
        # Keep only the highest version from each group
        deduplicated = []
        duplicates_removed = 0
        version_errors = 0
        
        for group in groups.values():
            if len(group) == 1:
                deduplicated.append(group[0])
            else:
                # Find the highest version
                try:
                    highest_version = max(group, key=lambda x: int(x.get('version', 0)))
                    deduplicated.append(highest_version)
                    duplicates_removed += len(group) - 1
                    
                    # Log duplicate removal for debugging
                    versions = [x.get('version', 0) for x in group]
                    self.logger.debug("Removed duplicates", extra={
                        'operation': 'deduplicate',
                        'report_id_hash': hashlib.sha256(
                            str(group[0].get('id', '')).encode()
                        ).hexdigest()[:8],
                        'versions_found': versions,
                        'kept_version': highest_version.get('version', 0),
                        'correlation_id': self.correlation_id
                    })
                    
                except (ValueError, TypeError) as e:
                    # If version comparison fails, take the first one
                    version_errors += 1
                    deduplicated.append(group[0])
                    self.logger.warning("Version comparison failed, using first report", extra={
                        'operation': 'deduplicate',
                        'error': str(e),
                        'group_size': len(group),
                        'correlation_id': self.correlation_id
                    })
        
        duration = time.time() - start_time
        self.logger.info("Deduplicated reports", extra={
            'operation': 'deduplicate',
            'input_count': len(reports),
            'output_count': len(deduplicated),
            'duplicates_removed': duplicates_removed,
            'invalid_reports': invalid_reports,
            'version_errors': version_errors,
            'duration': duration,
            'correlation_id': self.correlation_id
        })
        
        return deduplicated
    
    def truncate_summary(self, summary: Optional[str], word_limit: int = 300) -> str:
        """Truncate summary to specified word limit with safe handling.
        
        Args:
            summary: The summary text to truncate (may be None or empty)
            word_limit: Maximum number of words to retain
            
        Returns:
            Truncated summary string, empty string if input is invalid
            
        Raises:
            ValidationError: If word_limit is invalid
        """
        # Validate word limit
        if not isinstance(word_limit, int) or word_limit < 1:
            raise ValidationError("Word limit must be a positive integer", "word_limit", word_limit)
        
        # Handle None or empty summary
        if not summary or not isinstance(summary, str):
            return ""
        
        # Clean up the summary text - remove excessive whitespace and control characters
        cleaned_summary = re.sub(r'\s+', ' ', summary.strip())
        cleaned_summary = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', cleaned_summary)  # Remove control chars
        
        if not cleaned_summary:
            return ""
        
        words = cleaned_summary.split()
        if len(words) <= word_limit:
            return cleaned_summary
        
        # Truncate and add ellipsis
        truncated = ' '.join(words[:word_limit])
        return truncated + "..."
    
    def format_authors(self, authors: Optional[List[Any]]) -> str:
        """Format authors list into a readable string with validation.
        
        Args:
            authors: List of author dictionaries or strings
            
        Returns:
            Semicolon-separated string of author names
        """
        if not authors or not isinstance(authors, list):
            return ""
        
        author_names = []
        invalid_authors = 0
        
        for author in authors:
            author_name = None
            
            if isinstance(author, dict) and 'author' in author:
                author_name = author['author']
            elif isinstance(author, str):
                author_name = author
            else:
                invalid_authors += 1
                continue
            
            # Validate and clean author name
            if isinstance(author_name, str) and author_name.strip():
                # Clean up author name - remove excessive whitespace
                clean_name = re.sub(r'\s+', ' ', author_name.strip())
                # Remove potentially problematic characters
                clean_name = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', clean_name)
                
                if clean_name and len(clean_name) <= 200:  # Reasonable name length limit
                    author_names.append(clean_name)
                else:
                    invalid_authors += 1
            else:
                invalid_authors += 1
        
        if invalid_authors > 0:
            self.logger.debug("Skipped invalid authors", extra={
                'operation': 'format_authors',
                'invalid_count': invalid_authors,
                'valid_count': len(author_names),
                'correlation_id': self.correlation_id
            })
        
        return "; ".join(author_names)
    
    def format_date(self, date_str: Optional[str]) -> str:
        """Format date string for better readability in Excel.
        
        Args:
            date_str: ISO format date string to format
            
        Returns:
            Formatted date string (YYYY-MM-DD) or empty string if invalid
        """
        if not date_str or not isinstance(date_str, str):
            return ""
        
        date_str = date_str.strip()
        if not date_str:
            return ""
        
        try:
            # Handle various ISO format variations
            # Replace 'Z' with timezone offset for fromisoformat compatibility
            normalized_date = date_str.replace('Z', '+00:00')
            
            # Parse the date and reformat it
            date_obj = datetime.fromisoformat(normalized_date)
            return date_obj.strftime('%Y-%m-%d')
            
        except ValueError as e:
            # Log parsing failures for debugging
            self.logger.debug("Date parsing failed", extra={
                'operation': 'format_date',
                'date_length': len(date_str),
                'error': str(e),
                'correlation_id': self.correlation_id
            })
            
            # Try to extract just the date part if it's a longer string
            date_match = re.match(r'(\d{4}-\d{2}-\d{2})', date_str)
            if date_match:
                return date_match.group(1)
            
            # If all parsing fails, return empty string for safety
            return ""
    
    def create_word_document(self, reports_data: List[Dict[str, Any]], filename: str) -> None:
        """Create Word document with the processed data using professional formatting.
        
        Args:
            reports_data: List of processed report dictionaries
            filename: Output filename (must be .docx in current directory)
            
        Raises:
            ValidationError: If filename is invalid
            DataProcessingError: If Word document creation fails
        """
        # Validate inputs
        if not isinstance(reports_data, list):
            raise ValidationError("Reports data must be a list", "reports_data", type(reports_data))
        
        validated_path = InputValidator.validate_filename(filename)
        
        start_time = time.time()
        
        self.logger.info("Creating Word document", extra={
            'operation': 'create_document',
            'output_file': validated_path.name,
            'record_count': len(reports_data),
            'correlation_id': self.correlation_id
        })
        
        try:
            # Create a new Word document
            doc = Document()
            
            # Set up document title
            title = doc.add_heading('CRS Summary Report', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add generation date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            date_run = date_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
            date_run.italic = True
            
            # Add spacing
            doc.add_paragraph()
            
            successful_reports = 0
            failed_reports = 0
            
            for i, data in enumerate(reports_data):
                try:
                    if not isinstance(data, dict):
                        failed_reports += 1
                        continue
                    
                    # Extract and validate data
                    title = str(data.get('title', 'Untitled Report')).strip()
                    authors = self.format_authors(data.get('authors'))
                    date_published = self.format_date(data.get('publishDate'))
                    summary = self.truncate_summary(data.get('summary'))
                    report_id = str(data.get('id', 'Unknown ID')).strip()
                    url = str(data.get('url', '')).strip()
                    
                    # Add report title as Level 2 heading
                    if title:
                        title_heading = doc.add_heading(title, level=2)
                    else:
                        title_heading = doc.add_heading('Untitled Report', level=2)
                    
                    # Add author and date as subheading
                    subheading_parts = []
                    if authors:
                        subheading_parts.append(f"By {authors}")
                    if date_published:
                        subheading_parts.append(f"Published {date_published}")
                    
                    if subheading_parts:
                        subheading_para = doc.add_paragraph()
                        subheading_run = subheading_para.add_run(" | ".join(subheading_parts))
                        subheading_run.bold = True
                        subheading_run.font.size = Inches(0.12)  # Slightly smaller than normal
                    
                    # Add summary as body text
                    if summary:
                        summary_para = doc.add_paragraph()
                        summary_para.add_run(summary)
                    
                    # Add ID and URL in italics at the end
                    metadata_parts = []
                    if report_id:
                        metadata_parts.append(f"Report ID: {report_id}")
                    if url:
                        metadata_parts.append(f"URL: {url}")
                    
                    if metadata_parts:
                        metadata_para = doc.add_paragraph()
                        metadata_run = metadata_para.add_run(" | ".join(metadata_parts))
                        metadata_run.italic = True
                        metadata_run.font.size = Inches(0.1)  # Smaller font for metadata
                    
                    # Add one line break between reports (except for the last one)
                    if i < len(reports_data) - 1:
                        doc.add_paragraph()
                    
                    successful_reports += 1
                    
                except Exception as e:
                    failed_reports += 1
                    self.logger.warning("Failed to add report to document", extra={
                        'operation': 'create_document',
                        'report_index': i,
                        'error': str(e),
                        'correlation_id': self.correlation_id
                    })
            
            # Save the document
            doc.save(str(validated_path))
            
            # Verify file was created and has content
            if not validated_path.exists():
                raise DataProcessingError(f"Word document was not created: {validated_path}")
            
            file_size = validated_path.stat().st_size
            duration = time.time() - start_time
            
            self.logger.info("Word document created successfully", extra={
                'operation': 'create_document',
                'output_file': validated_path.name,
                'successful_reports': successful_reports,
                'failed_reports': failed_reports,
                'file_size_bytes': file_size,
                'duration': duration,
                'correlation_id': self.correlation_id
            })
            
            # User-facing success message (separate from logs)
            print(f"[SUCCESS] Word document '{validated_path.name}' created with {successful_reports} reports")
            if failed_reports > 0:
                print(f"[WARNING] {failed_reports} reports failed to process (see logs for details)")
                
        except OSError as e:
            duration = time.time() - start_time
            error_msg = f"Failed to create Word document: {e}"
            self.logger.error("File system error creating Word document", extra={
                'operation': 'create_document',
                'output_file': validated_path.name,
                'error': str(e),
                'duration': duration,
                'correlation_id': self.correlation_id
            })
            raise DataProcessingError(error_msg) from e
        
        except Exception as e:
            duration = time.time() - start_time
            error_msg = f"Unexpected error creating Word document: {e}"
            self.logger.error("Unexpected Word document creation error", extra={
                'operation': 'create_document',
                'error_type': type(e).__name__,
                'error': str(e),
                'duration': duration,
                'correlation_id': self.correlation_id
            })
            raise DataProcessingError(error_msg) from e
    
    def run(self, output_filename: Optional[str] = None) -> None:
        """Main execution method with comprehensive error handling and progress tracking.
        
        This method orchestrates the complete CRS data extraction workflow:
        1. Fetches list of CRS documents from Congress.gov API
        2. Filters for active documents only
        3. Removes duplicates based on ID and publish date
        4. Fetches detailed information for each document
        5. Creates formatted Word document output
        
        Args:
            output_filename: Name of the Word document file to create (if None, uses date-based filename)
            
        Raises:
            ConfigurationError: If system is not properly configured
            APIError: If API communication fails critically
            DataProcessingError: If data processing fails
        """
        # Generate filename with today's date if not provided
        if output_filename is None:
            today = datetime.now().strftime('%Y-%m-%d')
            output_filename = f'crs_summaries_{today}.docx'
        
        workflow_start_time = time.time()
        
        # User-facing messages (separate from structured logs)
        print("Starting CRS Summary Extraction...")
        print(f"Session ID: {self.correlation_id}")
        
        self.logger.info("Starting CRS extraction workflow", extra={
            'operation': 'workflow_start',
            'output_filename': output_filename,
            'correlation_id': self.correlation_id
        })
        
        try:
            # Step 1: Fetch list of CRS documents
            print("\n[1/5] Fetching CRS document list...")
            reports = self.fetch_crs_list()
            print(f"Found {len(reports)} total reports")
            
            if not reports:
                print("[ERROR] No reports found. Please check your API key and connection.")
                self.logger.error("No reports retrieved from API", extra={
                    'operation': 'workflow',
                    'step': 'fetch_list',
                    'correlation_id': self.correlation_id
                })
                return
            
            # Step 2: Filter for active reports only
            print("\n[2/5] Filtering for active reports...")
            active_reports = self.filter_active_reports(reports)
            print(f"Found {len(active_reports)} active reports")
            
            if not active_reports:
                print("[WARNING] No active reports found.")
                self.logger.warning("No active reports after filtering", extra={
                    'operation': 'workflow',
                    'step': 'filter_active',
                    'total_reports': len(reports),
                    'correlation_id': self.correlation_id
                })
                return
            
            # Step 3: Remove duplicates
            print("\n[3/5] Removing duplicates...")
            unique_reports = self.deduplicate_reports(active_reports)
            duplicates_removed = len(active_reports) - len(unique_reports)
            print(f"After deduplication: {len(unique_reports)} unique reports")
            if duplicates_removed > 0:
                print(f"Removed {duplicates_removed} duplicate reports")
            
            # Step 4: Fetch detailed information for each report
            print("\n[4/5] Fetching detailed information...")
            detailed_reports = []
            fetch_failures = 0
            
            total_reports = len(unique_reports)
            for i, report in enumerate(unique_reports, 1):
                report_id = report.get('id', 'unknown')
                
                # Progress indicator for user
                if i % 10 == 0 or i == total_reports:
                    print(f"Processing {i}/{total_reports} reports...")
                
                try:
                    detail = self.fetch_crs_detail(report_id)
                    if detail:
                        # Merge the list data with detailed data
                        merged_data = {**report, **detail}
                        detailed_reports.append(merged_data)
                    else:
                        fetch_failures += 1
                        
                except Exception as e:
                    fetch_failures += 1
                    self.logger.warning("Failed to fetch report details", extra={
                        'operation': 'workflow',
                        'step': 'fetch_details',
                        'report_id_hash': hashlib.sha256(
                            str(report_id).encode()
                        ).hexdigest()[:8],
                        'error': str(e),
                        'correlation_id': self.correlation_id
                    })
            
            success_rate = (len(detailed_reports) / total_reports) * 100 if total_reports > 0 else 0
            print(f"Successfully fetched details for {len(detailed_reports)} reports ({success_rate:.1f}% success rate)")
            
            if fetch_failures > 0:
                print(f"[WARNING] {fetch_failures} reports failed to fetch detailed information")
            
            if not detailed_reports:
                print("[ERROR] No detailed reports retrieved. Cannot create Word document.")
                self.logger.error("No detailed reports available for Word document creation", extra={
                    'operation': 'workflow',
                    'step': 'fetch_details',
                    'total_attempts': total_reports,
                    'failures': fetch_failures,
                    'correlation_id': self.correlation_id
                })
                return
            
            # Step 5: Create Word document
            print("\n[5/5] Creating Word document...")
            self.create_word_document(detailed_reports, output_filename)
            
            # Final summary
            workflow_duration = time.time() - workflow_start_time
            print(f"\n[SUCCESS] Process completed in {workflow_duration:.1f} seconds!")
            print(f"Created '{output_filename}' with {len(detailed_reports)} CRS reports")
            
            self.logger.info("CRS extraction workflow completed successfully", extra={
                'operation': 'workflow_complete',
                'total_duration': workflow_duration,
                'initial_reports': len(reports),
                'active_reports': len(active_reports),
                'unique_reports': len(unique_reports),
                'detailed_reports': len(detailed_reports),
                'fetch_failures': fetch_failures,
                'output_filename': output_filename,
                'correlation_id': self.correlation_id
            })
            
        except KeyboardInterrupt:
            print("\n[INTERRUPTED] Process interrupted by user")
            self.logger.info("Workflow interrupted by user", extra={
                'operation': 'workflow_interrupted',
                'correlation_id': self.correlation_id
            })
            raise
            
        except (ConfigurationError, APIError, DataProcessingError) as e:
            workflow_duration = time.time() - workflow_start_time
            print(f"\n[ERROR] {e}")
            self.logger.error("Workflow failed with known error", extra={
                'operation': 'workflow_failed',
                'error_type': type(e).__name__,
                'error': str(e),
                'duration': workflow_duration,
                'correlation_id': self.correlation_id
            })
            raise
            
        except Exception as e:
            workflow_duration = time.time() - workflow_start_time
            print(f"\n[ERROR] Unexpected error: {e}")
            self.logger.error("Workflow failed with unexpected error", extra={
                'operation': 'workflow_failed',
                'error_type': type(e).__name__,
                'error': str(e),
                'duration': workflow_duration,
                'correlation_id': self.correlation_id
            })
            raise DataProcessingError(f"Unexpected error during processing: {e}") from e

def main() -> None:
    """Main entry point with comprehensive error handling and user guidance.
    
    This function provides a user-friendly interface to the CRS extraction system
    with clear error messages and troubleshooting guidance.
    """
    import argparse
    
    # Set up command line arguments
    parser = argparse.ArgumentParser(
        description='Extract CRS summaries from Congress.gov API',
        epilog='Example: python crs_summary_extractor.py --output my_summaries.docx --verbose'
    )
    parser.add_argument(
        '--output', '-o',
        default=None,
        help='Output Word document filename (default: crs_summaries_YYYY-MM-DD.docx)'
    )
    parser.add_argument(
        '--verbose', '-v',
        action='store_true',
        help='Enable verbose logging'
    )
    parser.add_argument(
        '--json-logs',
        action='store_true',
        help='Output structured JSON logs'
    )
    
    args = parser.parse_args()
    
    try:
        # Initialize extractor with user preferences
        extractor = CRSSummaryExtractor(verbose=args.verbose, json_logs=args.json_logs)
        
        # Run the extraction process
        extractor.run(args.output)
        
        # Exit successfully
        sys.exit(0)
        
    except ConfigurationError as e:
        print(f"\n[CONFIGURATION ERROR] {e}")
        print("\nTroubleshooting:")
        print("1. Check that your .env file exists")
        print("2. Ensure it contains: CONGRESSGOV_API_KEY=your_actual_key")
        print("3. Verify your API key is valid and active")
        sys.exit(1)
        
    except APIError as e:
        print(f"\n[API ERROR] {e}")
        print("\nTroubleshooting:")
        print("1. Check your internet connection")
        print("2. Verify your API key is still valid")
        print("3. Try again in a few minutes (may be temporary API issue)")
        if hasattr(e, 'status_code') and e.status_code:
            print(f"4. HTTP Status Code: {e.status_code}")
        sys.exit(2)
        
    except DataProcessingError as e:
        print(f"\n[DATA PROCESSING ERROR] {e}")
        print("\nTroubleshooting:")
        print("1. Check available disk space")
        print("2. Ensure you have write permissions in this directory")
        print("3. Try a different output filename")
        sys.exit(3)
        
    except ValidationError as e:
        print(f"\n[VALIDATION ERROR] {e}")
        print("\nTroubleshooting:")
        print("1. Check your command line arguments")
        print("2. Ensure output filename ends with .docx")
        print("3. Verify all parameters are valid")
        sys.exit(4)
        
    except KeyboardInterrupt:
        print("\n\n[INTERRUPTED] Process stopped by user")
        print("You can restart the process at any time.")
        sys.exit(130)  # Standard exit code for Ctrl+C
        
    except Exception as e:
        print(f"\n[UNEXPECTED ERROR] {e}")
        print("\nThis is an unexpected error. Please report this issue with:")
        print(f"- Error message: {e}")
        print(f"- Error type: {type(e).__name__}")
        print("- Your Python version and operating system")
        
        # Log the full traceback for debugging
        import traceback
        logger = logging.getLogger('crs_extractor')
        logger.error("Unexpected error with traceback", extra={
            'operation': 'main_error',
            'error': str(e),
            'traceback': traceback.format_exc()
        })
        
        sys.exit(99)


if __name__ == "__main__":
    main()
